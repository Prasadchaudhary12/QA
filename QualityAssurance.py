import streamlit as st
import pandas as pd
import datetime
from io import BytesIO
from reportlab.pdfgen import canvas
from docx import Document

# -----------------------
# PAGE CONFIG + THEME
# -----------------------
st.set_page_config(page_title="Internal Audit QA Tool", layout="wide")

st.markdown("""
<style>
body, html {
    font-family: Calibri;
}
.stApp {
    background-color: #f4f4f4;
}
.stButton>button {
    background-color: #f1c40f;
    color: black;
    border-radius: 6px;
    font-weight: bold;
}
.sidebar .sidebar-content {
    background-color: #2b2b2b;
    color: white;
}
</style>
""", unsafe_allow_html=True)

# -----------------------
# SESSION STATE INIT
# -----------------------
def init():
    defaults = {
        "users": {"admin": "admin"},
        "logged_in": False,
        "user": "",
        "clients": [],
        "engagements": [],
        "qa_data": [],
        "logs": []
    }
    for k,v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init()

# -----------------------
# MASTER CHECKLIST
# -----------------------
MASTER_CHECKLIST = [
    "Audit Planning",
    "Risk Assessment",
    "Control Testing",
    "Evidence Verification",
    "Audit Conclusion"
]

MANDATORY_DOCS = [
    "Audit Scoping Memo",
    "Audit Report",
    "RCM",
    "Audit Program",
    "Workpapers",
    "Evidence"
]

# -----------------------
# LOG FUNCTION
# -----------------------
def log(action):
    st.session_state["logs"].append({
        "User": st.session_state["user"],
        "Action": action,
        "Time": datetime.datetime.now()
    })

# -----------------------
# LOGIN
# -----------------------
def login():
    st.title("🔐 Internal Audit QA Tool")

    u = st.text_input("Username")
    p = st.text_input("Password", type="password")

    if st.button("Login"):
        if u in st.session_state["users"] and st.session_state["users"][u] == p:
            st.session_state["logged_in"] = True
            st.session_state["user"] = u
            log("Login")
            st.rerun()
        else:
            st.error("Invalid credentials")

# -----------------------
# DASHBOARD
# -----------------------
def dashboard():
    st.title("📊 Dashboard")

    df = pd.DataFrame(st.session_state["qa_data"])

    total = len(df)
    completed = len(df[df["status"]=="Completed"]) if not df.empty else 0
    inprogress = len(df[df["status"]=="In Progress"]) if not df.empty else 0
    notstarted = total - completed - inprogress

    pass_count = len(df[df["result"]=="Pass"]) if not df.empty else 0
    fail_count = len(df[df["result"]=="Fail"]) if not df.empty else 0

    c1,c2,c3,c4 = st.columns(4)
    c1.metric("Total QA", total)
    c2.metric("Completed", completed)
    c3.metric("In Progress", f"{inprogress}")
    c4.metric("Not Started", notstarted)

    st.progress(completed/(total+1))

    st.write("✅ Pass:", pass_count, "❌ Fail:", fail_count)

# -----------------------
# CREATE CLIENT
# -----------------------
def create_client():
    st.title("🏢 Create Client")
    name = st.text_input("Client Name")

    if st.button("Add Client"):
        st.session_state["clients"].append(name)
        log(f"Client Created {name}")
        st.success("Client saved")

# -----------------------
# CREATE ENGAGEMENT
# -----------------------
def create_engagement():
    st.title("📁 Create Engagement")

    client = st.selectbox("Client", st.session_state["clients"])
    fy = st.text_input("Financial Year")
    process = st.text_input("Audit Process")
    auditor = st.text_input("Auditor Name")
    auditee = st.text_input("Auditee Name")
    dept = st.text_input("Department")
    title = st.text_input("Title")

    if st.button("Create"):
        st.session_state["engagements"].append({
            "client": client,
            "fy": fy,
            "process": process,
            "auditor": auditor,
            "auditee": auditee,
            "dept": dept,
            "title": title,
            "checklist": MASTER_CHECKLIST,
            "docs": {}
        })
        log("Engagement Created")
        st.success("Engagement created")

# -----------------------
# CHECKLIST
# -----------------------
def checklist():
    st.title("✅ QA Checklist")

    if not st.session_state["engagements"]:
        st.warning("No engagement created")
        return

    eng = st.selectbox("Select Engagement", st.session_state["engagements"])

    # Mandatory documents
    st.subheader("📎 Mandatory Uploads")
    for doc in MANDATORY_DOCS:
        eng["docs"][doc] = st.file_uploader(doc, key=doc)

    st.divider()

    for step in eng["checklist"]:
        st.subheader(step)

        doc = st.file_uploader("Upload Evidence", key=step)

        remark = st.text_area("Remarks", key=step+"r")

        col1,col2,col3 = st.columns(3)

        if col1.button("✔ Pass", key=step+"p"):
            save(step, "Pass", remark)

        if col2.button("❌ Fail", key=step+"f"):
            save(step, "Fail", remark)

        if col3.button("N/A", key=step+"na"):
            save(step, "NA", remark)

        # Chat
        if st.button("💬 Chat Assist", key=step+"c"):
            prompt = st.text_input("Refine step")
            st.info(f"AI Suggestion: Improve documentation for {step}")

def save(step,result,remark):
    st.session_state["qa_data"].append({
        "step": step,
        "result": result,
        "remark": remark,
        "status": "Completed"
    })
    log(f"{step} marked {result}")
    st.success("Saved ✔")

# -----------------------
# REPORT
# -----------------------
def report():
    st.title("📄 Final Report")

    df = pd.DataFrame(st.session_state["qa_data"])

    st.dataframe(df)

    # Excel
    st.download_button("Download Excel", df.to_csv(index=False), "QA.csv")

    # PDF
    def pdf_download():
        buffer = BytesIO()
        c = canvas.Canvas(buffer)
        c.drawString(100, 800, "QA Report")
        c.save()
        return buffer

    st.download_button("Download PDF", pdf_download(), "QA.pdf")

    # Word
    def word_download():
        doc = Document()
        doc.add_heading("QA Report", 0)
        for _,row in df.iterrows():
            doc.add_paragraph(str(row.to_dict()))
        buffer = BytesIO()
        doc.save(buffer)
        return buffer

    st.download_button("Download Word", word_download(), "QA.docx")

    # Chat refine report
    if st.button("💬 Refine Report"):
        st.info("AI Suggestion: Improve summary and observations.")

# -----------------------
# LOGS
# -----------------------
def logs():
    st.title("📜 Audit Logs")
    st.dataframe(pd.DataFrame(st.session_state["logs"]))

# -----------------------
# ARCHIVE
# -----------------------
def archive():
    st.title("📦 Archive")

    if st.button("Archive Data"):
        st.session_state["qa_data"] = []
        log("Archived")
        st.success("Archived")

# -----------------------
# MAIN
# -----------------------
if not st.session_state["logged_in"]:
    login()
else:
    st.sidebar.title("🚆 QA System")

    st.sidebar.write(f"👤 {st.session_state['user']}")

    if st.sidebar.button("🚪 Logout"):
        log("Logout")
        st.session_state["logged_in"] = False
        st.session_state["user"] = ""
        st.rerun()

    menu = st.sidebar.radio("Menu", [
        "Home",
        "Dashboard",
        "Create Client",
        "Create Engagement",
        "Checklist",
        "Report",
        "Logs",
        "Archive"
    ])

    if menu == "Home":
        st.title("🏠 Welcome to QA Tool")
        st.write("System operates similar to structured workflow platforms.")

    elif menu == "Dashboard":
        dashboard()

    elif menu == "Create Client":
        create_client()

    elif menu == "Create Engagement":
        create_engagement()

    elif menu == "Checklist":
        checklist()

    elif menu == "Report":
        report()

    elif menu == "Logs":
        logs()

    elif menu == "Archive":
        archive()
``
